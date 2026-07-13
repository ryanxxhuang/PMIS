from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parent


def font_name():
    candidates = [
        Path('/System/Library/Fonts/PingFang.ttc'),
        Path('/System/Library/Fonts/Supplemental/Arial Unicode.ttf'),
    ]
    for path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont('CJK', str(path), subfontIndex=0))
            return 'CJK'
    return 'Helvetica'


FONT = font_name()


def make_pdf(path, pages):
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    for page_no, lines in enumerate(pages, 1):
        c.setFont(FONT, 12)
        y = height - 64
        c.drawString(60, y, f'PMIS 第四輪測試文件｜第 {page_no} 頁')
        y -= 36
        for line in lines:
            c.drawString(60, y, line)
            y -= 24
        c.showPage()
    c.save()


make_pdf(ROOT / 'R4-施工契約-可擷取.pdf', [
    [
        '公共工程施工契約（第四輪驗收專用）',
        '第 5 條：廠商應於接獲開工通知次日起 7 日內申報開工。',
        '逾期每一日違約金新臺幣 10,000 元。',
        '原文識別：R4-CONTRACT-P1-START-7D。',
    ],
    [
        '第 12 條：廠商應於每月 5 日前提送前一月份施工月報。',
        '應留存施工月報、監造審查意見與機關核定函。',
        '原文識別：R4-CONTRACT-P2-MONTHLY-REPORT。',
    ],
    [
        '第 18 條：竣工後 15 日內提交竣工圖、結算明細及試驗報告。',
        '未依期限提送者，每逾一日扣款新臺幣 5,000 元。',
        '原文識別：R4-CONTRACT-P3-CLOSEOUT-15D。',
    ],
])

make_pdf(ROOT / 'R4-非契約-員工餐廳菜單.pdf', [[
    '員工餐廳七月份菜單',
    '星期一：咖哩飯。星期二：牛肉麵。星期三：素食便當。',
    '本文件不含工程契約、履約期限、罰則或施工要求。',
    '原文識別：R4-NON-CONTRACT-MENU。',
]])

doc = Document()
doc.add_heading('公共工程補充契約（第四輪 Word 驗收）', 0)
doc.add_paragraph('第 21 條：承攬廠商應於材料進場前 10 日提送材料送審文件，經監造核備後始得使用。')
doc.add_paragraph('應留存材料型錄、試驗報告及核備紀錄。原文識別：R4-DOCX-MATERIAL-10D。')
doc.add_paragraph('第 22 條：缺失通知後 3 日內提出改善計畫。逾期每日扣款新臺幣 2,000 元。')
doc.add_paragraph('DOCX 無可靠頁界，擷取來源不得偽造頁碼。')
doc.save(ROOT / 'R4-補充契約-可擷取.docx')

(ROOT / 'R4-壞檔.pdf').write_text('THIS IS NOT A PDF. R4-CORRUPT-PDF', encoding='utf-8')
(ROOT / 'R4-壞檔.docx').write_text('THIS IS NOT A DOCX. R4-CORRUPT-DOCX', encoding='utf-8')
(ROOT / 'R4-壞XML.xml').write_text('<ETenderSheet><broken>', encoding='utf-8')
(ROOT / 'R4-非PCCES.xml').write_text(
    '<?xml version="1.0"?><employees><employee id="1">R4-NON-PCCES</employee></employees>',
    encoding='utf-8',
)

source = (ROOT.parent / 'qa-e2e' / 'boq.xml').read_bytes()
close = source.rfind(b'</ETenderSheet>')
padding = b'<!--' + (b'R4-LARGE-PCCES-PADDING-' * 900000) + b'-->'
(ROOT / 'R4-超大PCCES.xml').write_bytes(source[:close] + padding + source[close:])

print('fixtures generated')
for path in sorted(ROOT.iterdir()):
    if path.is_file() and path.name not in {'generate_fixtures.py', 'setup.mjs'}:
        print(f'{path.name}\t{path.stat().st_size}')
